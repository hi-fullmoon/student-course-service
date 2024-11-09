import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches
from fastapi.responses import StreamingResponse
from typing import List
from ..models import models

def create_excel_report(data: List[dict], filename: str) -> StreamingResponse:
    df = pd.DataFrame(data)
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    output.seek(0)

    headers = {
        'Content-Disposition': f'attachment; filename="{filename}"'
    }
    return StreamingResponse(output, media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', headers=headers)

def create_word_report(title: str, content: List[dict], filename: str) -> StreamingResponse:
    doc = Document()
    doc.add_heading(title, 0)

    for item in content:
        for key, value in item.items():
            doc.add_paragraph(f"{key}: {value}")
        doc.add_paragraph("-------------------")

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    headers = {
        'Content-Disposition': f'attachment; filename="{filename}"'
    }
    return StreamingResponse(output, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers=headers)